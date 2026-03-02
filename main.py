import pandas as pd
import os
from docx import Document
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
from collections import defaultdict
import datetime
import json
import platform

class SettingsManager:
    """Класс для управления настройками"""
    def __init__(self):
        self.config_dir = os.path.expanduser("~/.course_work_generator")
        self.config_file = os.path.join(self.config_dir, "settings.json")
        self.load_settings()
    
    def load_settings(self):
        """Загрузка настроек из файла"""
        default_settings = {
            "last_discipline": "",
            "last_excel_path": "",
            "last_title_template": "",
            "last_task_template": "",
            "last_review_template": "",
            "last_output_dir": os.path.expanduser("~/Desktop/документы курсовой работы"),
            "show_details": True,
            "auto_open_folder": False
        }
        
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                    # Обновляем только существующие ключи
                    for key in default_settings:
                        if key not in settings:
                            settings[key] = default_settings[key]
                    self.settings = settings
            else:
                # Создаем папку для конфига, если её нет
                os.makedirs(self.config_dir, exist_ok=True)
                self.settings = default_settings
                self.save_settings()
        except Exception as e:
            print(f"Ошибка загрузки настроек: {e}")
            self.settings = default_settings
    
    def save_settings(self):
        """Сохранение настроек в файл"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Ошибка сохранения настроек: {e}")
    
    def get(self, key, default=None):
        return self.settings.get(key, default)
    
    def set(self, key, value):
        self.settings[key] = value
        self.save_settings()

class DocumentGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор документов курсовой работы")
        self.root.geometry("1200x800")
        
        # Загружаем настройки
        self.settings = SettingsManager()
        
        # Переменные
        self.excel_path = tk.StringVar(value=self.settings.get("last_excel_path", ""))
        self.discipline = tk.StringVar(value=self.settings.get("last_discipline", ""))
        self.title_template = tk.StringVar(value=self.settings.get("last_title_template", ""))
        self.task_template = tk.StringVar(value=self.settings.get("last_task_template", ""))
        self.review_template = tk.StringVar(value=self.settings.get("last_review_template", ""))
        self.output_dir = tk.StringVar(value=self.settings.get("last_output_dir", ""))
        self.df = None
        self.filtered_df = None
        self.selected_students = []
        
        # Переменные для фильтрации
        self.filter_text = tk.StringVar()
        self.filter_text.trace('w', lambda *args: self.apply_filter())
        
        self.filter_group = tk.StringVar()
        self.filter_supervisor = tk.StringVar()
        
        self.setup_ui()
        self.setup_hotkeys()
        
    def setup_hotkeys(self):
        """Настройка горячих клавиш"""
        self.root.bind('<Control-a>', lambda e: self.select_all())
        self.root.bind('<Control-d>', lambda e: self.deselect_all())
        self.root.bind('<Control-g>', lambda e: self.generate_documents())
        self.root.bind('<Control-o>', lambda e: self.open_output_folder())
        self.root.bind('<Control-f>', lambda e: self.filter_entry.focus())
        self.root.bind('<Control-s>', lambda e: self.open_settings())
        self.root.bind('<F1>', lambda e: self.show_help())
        
    def setup_ui(self):
        # Создаем canvas с прокруткой
        canvas = tk.Canvas(self.root)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Упаковываем canvas и scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Основной фрейм
        main_frame = ttk.Frame(self.scrollable_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Настройка весов для растягивания
        main_frame.columnconfigure(1, weight=1)
        
        row = 0
        
        # Верхняя панель с заголовком и кнопкой настроек
        top_frame = ttk.Frame(main_frame)
        top_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        top_frame.columnconfigure(0, weight=1)
        
        ttk.Label(top_frame, text="Генератор документов курсовой работы", 
                 font=('Arial', 16, 'bold')).grid(row=0, column=0, sticky=tk.W)
        
        ttk.Button(top_frame, text="⚙ Настройки", 
                  command=self.open_settings).grid(row=0, column=1, padx=5)
        
        ttk.Button(top_frame, text="❓ Помощь (F1)", 
                  command=self.show_help).grid(row=0, column=2, padx=5)
        row += 1
        
        # Дисциплина
        ttk.Label(main_frame, text="Название дисциплины:").grid(row=row, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.discipline, width=70).grid(row=row, column=1, sticky=(tk.W, tk.E), padx=5)
        row += 1
        
        # Excel файл
        ttk.Label(main_frame, text="Excel файл с данными:").grid(row=row, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.excel_path, width=70).grid(row=row, column=1, sticky=(tk.W, tk.E), padx=5)
        ttk.Button(main_frame, text="Обзор...", command=self.browse_excel).grid(row=row, column=2, padx=5)
        row += 1
        
        # Шаблоны
        templates = [
            ("Титульный лист:", self.title_template),
            ("Задание:", self.task_template),
            ("Отзыв:", self.review_template)
        ]
        
        for label, var in templates:
            ttk.Label(main_frame, text=label).grid(row=row, column=0, sticky=tk.W, pady=5)
            ttk.Entry(main_frame, textvariable=var, width=70).grid(row=row, column=1, sticky=(tk.W, tk.E), padx=5)
            ttk.Button(main_frame, text="Обзор...", 
                      command=lambda v=var: self.browse_file(v)).grid(row=row, column=2, padx=5)
            row += 1
        
        # Папка для сохранения
        ttk.Label(main_frame, text="Папка для сохранения:").grid(row=row, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.output_dir, width=70).grid(row=row, column=1, sticky=(tk.W, tk.E), padx=5)
        ttk.Button(main_frame, text="Обзор...", command=self.browse_output_dir).grid(row=row, column=2, padx=5)
        row += 1
        
        # Кнопка загрузки данных
        ttk.Button(main_frame, text="Загрузить данные из Excel (Ctrl+L)", 
                  command=self.load_excel_data).grid(row=row, column=0, columnspan=3, pady=10)
        row += 1
        
        # Панель фильтрации
        filter_frame = ttk.LabelFrame(main_frame, text="Фильтрация", padding="5")
        filter_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        filter_frame.columnconfigure(1, weight=1)
        
        ttk.Label(filter_frame, text="Поиск по ФИО (Ctrl+F):").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.filter_entry = ttk.Entry(filter_frame, textvariable=self.filter_text, width=30)
        self.filter_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=5)
        
        ttk.Label(filter_frame, text="Группа:").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.filter_group_combo = ttk.Combobox(filter_frame, textvariable=self.filter_group, width=15)
        self.filter_group_combo.grid(row=0, column=3, sticky=tk.W, padx=5)
        self.filter_group_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_filter())
        
        ttk.Label(filter_frame, text="Руководитель:").grid(row=0, column=4, sticky=tk.W, padx=5)
        self.filter_supervisor_combo = ttk.Combobox(filter_frame, textvariable=self.filter_supervisor, width=15)
        self.filter_supervisor_combo.grid(row=0, column=5, sticky=tk.W, padx=5)
        self.filter_supervisor_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_filter())
        
        ttk.Button(filter_frame, text="Сбросить фильтры", 
                  command=self.reset_filters).grid(row=0, column=6, padx=5)
        row += 1
        
        # Фрейм для выбора студентов
        self.selection_frame = ttk.LabelFrame(main_frame, text="Выберите студентов для генерации", padding="5")
        self.selection_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        self.selection_frame.columnconfigure(0, weight=1)
        self.selection_frame.rowconfigure(0, weight=1)
        row += 1
        
        # Создаем Treeview для таблицы студентов
        columns = ('Выбор', 'ФИО', 'Группа', 'Руководитель')
        self.tree = ttk.Treeview(self.selection_frame, columns=columns, show='headings', height=15)
        
        # Настраиваем заголовки
        self.tree.heading('Выбор', text='✓')
        self.tree.heading('ФИО', text='ФИО студента')
        self.tree.heading('Группа', text='Группа')
        self.tree.heading('Руководитель', text='Руководитель')
        
        # Настраиваем ширину колонок
        self.tree.column('Выбор', width=40, anchor='center')
        self.tree.column('ФИО', width=400)
        self.tree.column('Группа', width=120)
        self.tree.column('Руководитель', width=200)
        
        # Добавляем скроллбары
        vsb = ttk.Scrollbar(self.selection_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(self.selection_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Размещаем Treeview и скроллбары
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        vsb.grid(row=0, column=1, sticky=(tk.N, tk.S))
        hsb.grid(row=1, column=0, sticky=(tk.W, tk.E))
        
        # Привязываем обработчик клика
        self.tree.bind('<ButtonRelease-1>', self.on_tree_click)
        
        # Кнопки выбора
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=row, column=0, columnspan=3, pady=5)
        ttk.Button(button_frame, text="Выбрать всех (Ctrl+A)", command=self.select_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Снять все (Ctrl+D)", command=self.deselect_all).pack(side=tk.LEFT, padx=5)
        row += 1
        
        # Фрейм с прогрессом и деталями
        progress_frame = ttk.LabelFrame(main_frame, text="Прогресс генерации", padding="5")
        progress_frame.grid(row=row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        progress_frame.columnconfigure(1, weight=1)
        
        # Прогресс-бар
        ttk.Label(progress_frame, text="Общий прогресс:").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.progress = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=5)
        
        # Детальная информация
        self.details_var = tk.StringVar(value="Готов к работе")
        self.details_label = ttk.Label(progress_frame, textvariable=self.details_var, font=('Arial', 9))
        self.details_label.grid(row=1, column=0, columnspan=2, sticky=tk.W, padx=5, pady=2)
        
        # Статистика
        self.stats_var = tk.StringVar(value="")
        self.stats_label = ttk.Label(progress_frame, textvariable=self.stats_var, font=('Arial', 9))
        self.stats_label.grid(row=2, column=0, columnspan=2, sticky=tk.W, padx=5, pady=2)
        row += 1
        
        # Кнопка генерации
        self.generate_btn = ttk.Button(main_frame, text="Сгенерировать документы (Ctrl+G)", 
                                       command=self.generate_documents)
        self.generate_btn.grid(row=row, column=0, columnspan=3, pady=10)
        row += 1
        
        # Информация о пути сохранения
        self.path_info_var = tk.StringVar()
        ttk.Label(main_frame, textvariable=self.path_info_var, foreground="blue").grid(row=row, column=0, columnspan=3, pady=5)
        row += 1
        
        # Кнопка для открытия папки
        ttk.Button(main_frame, text="📂 Открыть папку с документами (Ctrl+O)", 
                  command=self.open_output_folder).grid(row=row, column=0, columnspan=3, pady=5)
        
        # Словарь для хранения состояния выбора
        self.selection_state = {}
        
        # Обновляем информацию о пути
        self.update_path_info()
    
    def open_settings(self):
        """Открыть окно настроек"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Настройки")
        settings_window.geometry("500x400")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        # Центрируем окно
        settings_window.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - settings_window.winfo_width()) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - settings_window.winfo_height()) // 2
        settings_window.geometry(f"+{x}+{y}")
        
        # Фрейм для настроек
        frame = ttk.Frame(settings_window, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)
        
        row = 0
        
        # Показывать детали
        show_details_var = tk.BooleanVar(value=self.settings.get("show_details", True))
        ttk.Checkbutton(frame, text="Показывать детальную информацию о прогрессе", 
                       variable=show_details_var).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=5)
        row += 1
        
        # Автоматически открывать папку
        auto_open_var = tk.BooleanVar(value=self.settings.get("auto_open_folder", False))
        ttk.Checkbutton(frame, text="Автоматически открывать папку после генерации", 
                       variable=auto_open_var).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=5)
        row += 1
        
        # Разделитель
        ttk.Separator(frame, orient='horizontal').grid(row=row, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10)
        row += 1
        
        # Информация о горячих клавишах
        ttk.Label(frame, text="Горячие клавиши:", font=('Arial', 10, 'bold')).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=5)
        row += 1
        
        shortcuts = [
            "Ctrl+A - Выбрать всех",
            "Ctrl+D - Снять всех",
            "Ctrl+G - Начать генерацию",
            "Ctrl+O - Открыть папку",
            "Ctrl+F - Фокус на поиск",
            "Ctrl+S - Открыть настройки",
            "F1 - Помощь"
        ]
        
        for shortcut in shortcuts:
            ttk.Label(frame, text=f"  • {shortcut}").grid(row=row, column=0, columnspan=2, sticky=tk.W)
            row += 1
        
        # Кнопки
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=row, column=0, columnspan=2, pady=20)
        
        def save_settings():
            # Сохраняем настройки
            self.settings.set("show_details", show_details_var.get())
            self.settings.set("auto_open_folder", auto_open_var.get())
            settings_window.destroy()
            messagebox.showinfo("Настройки", "Настройки сохранены")
        
        ttk.Button(button_frame, text="Сохранить", command=save_settings).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Отмена", command=settings_window.destroy).pack(side=tk.LEFT, padx=5)
    
    def show_help(self):
        """Показать окно помощи"""
        help_text = """
        Генератор документов курсовой работы
        ===================================
        
        Как использовать:
        1. Введите название дисциплины
        2. Укажите путь к Excel файлу с данными
        3. Укажите пути к шаблонам документов
        4. Выберите папку для сохранения
        5. Загрузите данные из Excel
        6. Отметьте нужных студентов
        7. Нажмите "Сгенерировать"
        
        Горячие клавиши:
        • Ctrl+A - Выбрать всех
        • Ctrl+D - Снять всех
        • Ctrl+G - Начать генерацию
        • Ctrl+O - Открыть папку
        • Ctrl+F - Фокус на поиск
        • Ctrl+S - Открыть настройки
        • F1 - Эта помощь
        
        Формат Excel файла:
        Необходимые колонки:
        - Студент ФИО
        - Студент инициалы
        - Группа
        - РУКОВОДИТЕЛЬ
        - Тема курсовой работы
        - ДАТА1 ... ДАТА9 (опционально)
        
        Поддерживаемые плейсхолдеры в шаблонах:
        __СТУДЕНТ__, __ИНИЦИАЛЫСТ__, __ГРУППА__,
        __РУКОВОДИТЕЛЬ__, __ТЕМА__, __ДОЛЖНОСТЬ__,
        __Название дисциплины__, __ДАТА1__ ... __ДАТА9__
        """
        
        messagebox.showinfo("Помощь", help_text)
    
    def browse_excel(self):
        filename = filedialog.askopenfilename(
            title="Выберите Excel файл",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if filename:
            self.excel_path.set(filename)
            self.settings.set("last_excel_path", filename)
    
    def browse_file(self, var):
        filename = filedialog.askopenfilename(
            title="Выберите файл шаблона",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            var.set(filename)
            # Сохраняем в настройках в зависимости от переменной
            if var == self.title_template:
                self.settings.set("last_title_template", filename)
            elif var == self.task_template:
                self.settings.set("last_task_template", filename)
            elif var == self.review_template:
                self.settings.set("last_review_template", filename)
    
    def browse_output_dir(self):
        directory = filedialog.askdirectory(
            title="Выберите папку для сохранения документов"
        )
        if directory:
            self.output_dir.set(directory)
            self.settings.set("last_output_dir", directory)
            self.update_path_info()
    
    def update_path_info(self):
        """Обновляет информацию о пути сохранения"""
        self.path_info_var.set(f"Документы будут сохранены в: {self.output_dir.get()}")
    
    def open_output_folder(self):
        """Открывает папку с документами в Finder"""
        folder = self.output_dir.get()
        if os.path.exists(folder):
            if platform.system() == 'Darwin':  # macOS
                os.system(f'open "{folder}"')
            elif platform.system() == 'Windows':
                os.system(f'explorer "{folder}"')
            else:  # Linux
                os.system(f'xdg-open "{folder}"')
        else:
            messagebox.showinfo("Информация", "Папка еще не создана. Сначала сгенерируйте документы.")
    
    def load_excel_data(self):
        if not self.excel_path.get():
            messagebox.showerror("Ошибка", "Выберите Excel файл")
            return
        
        try:
            self.df = pd.read_excel(self.excel_path.get())
            self.filtered_df = self.df.copy()
            
            # Сохраняем дисциплину
            self.settings.set("last_discipline", self.discipline.get())
            
            # Обновляем списки для фильтров
            groups = sorted(self.df['Группа'].dropna().unique())
            supervisors = sorted(self.df['РУКОВОДИТЕЛЬ'].dropna().unique())
            
            self.filter_group_combo['values'] = ['Все'] + list(groups)
            self.filter_supervisor_combo['values'] = ['Все'] + list(supervisors)
            
            self.reset_filters()
            
            messagebox.showinfo("Успех", f"Загружено {len(self.df)} записей")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить файл: {str(e)}")
    
    def apply_filter(self):
        """Применить фильтры к таблице"""
        if self.df is None:
            return
        
        # Очищаем текущее отображение
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Применяем фильтры
        filtered = self.df.copy()
        
        # Текстовый поиск по ФИО
        search_text = self.filter_text.get().lower()
        if search_text:
            filtered = filtered[filtered['Студент ФИО'].str.lower().str.contains(search_text, na=False)]
        
        # Фильтр по группе
        group = self.filter_group.get()
        if group and group != 'Все':
            filtered = filtered[filtered['Группа'] == group]
        
        # Фильтр по руководителю
        supervisor = self.filter_supervisor.get()
        if supervisor and supervisor != 'Все':
            filtered = filtered[filtered['РУКОВОДИТЕЛЬ'] == supervisor]
        
        self.filtered_df = filtered
        
        # Заполняем таблицу
        for idx, row in filtered.iterrows():
            fio = str(row.get('Студент ФИО', 'N/A'))
            group = str(row.get('Группа', 'N/A'))
            supervisor = str(row.get('РУКОВОДИТЕЛЬ', 'N/A'))
            
            # Сохраняем состояние выбора для отфильтрованных строк
            if idx not in self.selection_state:
                self.selection_state[idx] = False
            
            # Добавляем строку
            checkbox = '☒' if self.selection_state.get(idx, False) else '☐'
            self.tree.insert('', 'end', iid=str(idx), values=(checkbox, fio, group, supervisor))
    
    def reset_filters(self):
        """Сбросить все фильтры"""
        self.filter_text.set("")
        self.filter_group.set("")
        self.filter_supervisor.set("")
        self.apply_filter()
    
    def on_tree_click(self, event):
        """Обработчик клика по таблице"""
        region = self.tree.identify_region(event.x, event.y)
        if region == "cell":
            column = self.tree.identify_column(event.x)
            if column == '#1':  # Клик по колонке выбора
                item = self.tree.identify_row(event.y)
                if item:
                    idx = int(item)
                    # Инвертируем состояние
                    self.selection_state[idx] = not self.selection_state[idx]
                    # Обновляем отображение
                    values = list(self.tree.item(item, 'values'))
                    values[0] = '☒' if self.selection_state[idx] else '☐'
                    self.tree.item(item, values=values)
    
    def select_all(self):
        """Выбрать всех студентов (только из отфильтрованных)"""
        for item in self.tree.get_children():
            idx = int(item)
            self.selection_state[idx] = True
            values = list(self.tree.item(item, 'values'))
            values[0] = '☒'
            self.tree.item(item, values=values)
    
    def deselect_all(self):
        """Снять выбор со всех студентов (только из отфильтрованных)"""
        for item in self.tree.get_children():
            idx = int(item)
            self.selection_state[idx] = False
            values = list(self.tree.item(item, 'values'))
            values[0] = '☐'
            self.tree.item(item, values=values)
    
    def get_selected_indices(self):
        """Получить индексы выбранных студентов"""
        return [idx for idx, selected in self.selection_state.items() if selected]
    
    def check_templates(self):
        """Проверка существования файлов шаблонов"""
        templates = {
            "Титульный лист": self.title_template.get(),
            "Задание": self.task_template.get(),
            "Отзыв": self.review_template.get()
        }
        
        missing = []
        for name, path in templates.items():
            if not path:
                missing.append(f"{name} (путь не указан)")
            elif not os.path.exists(path):
                missing.append(f"{name} ({path})")
        
        if missing:
            messagebox.showerror("Ошибка", f"Отсутствуют файлы шаблонов:\n" + "\n".join(missing))
            return False
        return True
    
    def replace_text(self, doc_obj, replacements):
        """Функция замены плейсхолдеров"""
        for paragraph in doc_obj.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    full_text = "".join(run.text for run in paragraph.runs)
                    if key in full_text:
                        new_text = full_text.replace(key, str(value) if value else "")
                        for run in paragraph.runs:
                            run.text = ""
                        paragraph.runs[0].text = new_text
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_text(cell, replacements)
    
    def generate_documents(self):
        if not self.discipline.get():
            messagebox.showerror("Ошибка", "Введите название дисциплины")
            return
        
        if not self.check_templates():
            return
        
        if self.df is None:
            messagebox.showerror("Ошибка", "Сначала загрузите данные из Excel")
            return
        
        selected_indices = self.get_selected_indices()
        if not selected_indices:
            messagebox.showerror("Ошибка", "Выберите хотя бы одного студента")
            return
        
        # Проверяем, можно ли создать папку для сохранения
        try:
            os.makedirs(self.output_dir.get(), exist_ok=True)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать папку для сохранения:\n{str(e)}")
            return
        
        # Блокируем кнопку генерации
        self.generate_btn.config(state='disabled')
        
        # Запускаем генерацию в отдельном потоке
        thread = threading.Thread(target=self._generate_thread, args=(selected_indices,))
        thread.daemon = True
        thread.start()
    
    def _generate_thread(self, selected_indices):
        start_time = datetime.datetime.now()
        generated_count = 0
        errors_count = 0
        
        try:
            total = len(selected_indices)
            self.progress['maximum'] = total
            
            # Группируем студентов по группам
            students_by_group = defaultdict(list)
            for idx in selected_indices:
                row = self.df.loc[idx]
                group = str(row.get("Группа", "Без группы")).strip()
                students_by_group[group].append(idx)
            
            # Создаем базовую структуру папок
            for group in students_by_group.keys():
                group_dir = os.path.join(self.output_dir.get(), group)
                os.makedirs(os.path.join(group_dir, "задания"), exist_ok=True)
                os.makedirs(os.path.join(group_dir, "титулы"), exist_ok=True)
                os.makedirs(os.path.join(group_dir, "отзывы"), exist_ok=True)
            
            for progress_idx, idx in enumerate(selected_indices, 1):
                row = self.df.loc[idx]
                student_name = row.get('Студент ФИО', 'N/A')
                
                # Обновляем статус
                if self.settings.get("show_details", True):
                    self.details_var.set(f"Генерация: {student_name} ({progress_idx}/{total})")
                else:
                    self.details_var.set(f"Генерация документов...")
                
                # Извлечение данных
                группа = str(row.get("Группа", "Без группы")).strip()
                студент_фио = str(row.get("Студент ФИО", ""))
                студент_инициалы = str(row.get("Студент инициалы", "")) if pd.notna(row.get("Студент инициалы")) else ""
                
                # Формируем имя файла
                if not студент_инициалы and студент_фио:
                    parts = студент_фио.split()
                    if len(parts) >= 2:
                        фамилия = parts[0]
                        имя = parts[1][0] if len(parts[1]) > 0 else ""
                        отчество = parts[2][0] if len(parts) > 2 else ""
                        студент_инициалы = f"{фамилия} {имя}.{отчество}." if отчество else f"{фамилия} {имя}."
                
                safe_name = "".join(c for c in студент_инициалы if c.isalnum() or c in " .-").strip()
                if not safe_name:
                    safe_name = f"студент_{idx}"
                
                # Пути к папкам
                group_dir = os.path.join(self.output_dir.get(), группа)
                tasks_dir = os.path.join(group_dir, "задания")
                titles_dir = os.path.join(group_dir, "титулы")
                reviews_dir = os.path.join(group_dir, "отзывы")
                
                # Общие замены
                должность = str(row.get("ДОЛЖНОСТЬ", "")) if pd.notna(row.get("ДОЛЖНОСТЬ", "")) else ""
                
                replacements = {
                    "__СТУДЕНТ__": студент_фио,
                    "__ИНИЦИАЛЫСТ__": студент_инициалы,
                    "__ГРУППА__": группа,
                    "__РУКОВОДИТЕЛЬ__": str(row.get("РУКОВОДИТЕЛЬ", "")),
                    "__ТЕМА__": str(row.get("Тема курсовой работы", "")),
                    "__ДОЛЖНОСТЬ__": должность,
                    "__Название дисциплины__": self.discipline.get(),
                }
                
                # Добавляем даты
                for i in range(1, 10):
                    date_key = f"__ДАТА{i}__"
                    replacements[date_key] = str(row.get(f"ДАТА{i}", "")) if pd.notna(row.get(f"ДАТА{i}", "")) else ""
                
                # Генерация документов
                try:
                    # Титульный лист
                    title_doc = Document(self.title_template.get())
                    self.replace_text(title_doc, replacements)
                    title_path = os.path.join(titles_dir, f"титул_{safe_name}.docx")
                    title_doc.save(title_path)
                    
                    # Задание
                    task_doc = Document(self.task_template.get())
                    self.replace_text(task_doc, replacements)
                    task_path = os.path.join(tasks_dir, f"задание_{safe_name}.docx")
                    task_doc.save(task_path)
                    
                    # Отзыв
                    review_doc = Document(self.review_template.get())
                    self.replace_text(review_doc, replacements)
                    review_path = os.path.join(reviews_dir, f"отзыв_{safe_name}.docx")
                    review_doc.save(review_path)
                    
                    generated_count += 3
                    
                except Exception as e:
                    errors_count += 1
                    print(f"Ошибка при генерации для {студент_инициалы}: {str(e)}")
                    self.details_var.set(f"Ошибка при генерации для {студент_инициалы}")
                
                # Обновляем статистику
                elapsed_time = datetime.datetime.now() - start_time
                if progress_idx > 0:
                    avg_time = elapsed_time / progress_idx
                    remaining = avg_time * (total - progress_idx)
                    remaining_str = str(remaining).split('.')[0]
                else:
                    remaining_str = "расчет..."
                
                self.stats_var.set(f"Создано: {generated_count} файлов | Ошибок: {errors_count} | Осталось: ~{remaining_str}")
                
                # Обновляем прогресс
                self.progress['value'] = progress_idx
                self.root.update_idletasks()
            
            # Сохраняем настройки
            self.settings.set("last_discipline", self.discipline.get())
            
            # Итоговое сообщение
            total_time = datetime.datetime.now() - start_time
            total_time_str = str(total_time).split('.')[0]
            
            self.details_var.set(f"Генерация завершена за {total_time_str}")
            self.stats_var.set(f"Всего создано: {generated_count} файлов | Ошибок: {errors_count}")
            
            messagebox.showinfo("Успех", 
                              f"Генерация завершена!\n"
                              f"Создано файлов: {generated_count}\n"
                              f"Ошибок: {errors_count}\n"
                              f"Время: {total_time_str}\n\n"
                              f"Сохранено в: {self.output_dir.get()}")
            
            # Автоматически открываем папку, если включено в настройках
            if self.settings.get("auto_open_folder", False) and generated_count > 0:
                self.open_output_folder()
            
        except Exception as e:
            self.details_var.set(f"Ошибка: {str(e)}")
            messagebox.showerror("Ошибка", f"Произошла ошибка:\n{str(e)}")
        finally:
            self.progress['value'] = 0
            self.generate_btn.config(state='normal')
            self.update_path_info()

if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentGeneratorApp(root)
    root.mainloop()